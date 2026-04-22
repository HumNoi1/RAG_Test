from io import BytesIO

from docx import Document
from pypdf import PdfReader


class UnsupportedFileTypeError(ValueError):
    """Raised when the uploaded file type is not supported."""


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("tis-620", errors="replace")


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()


def _iter_docx_text(document: Document):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                yield " | ".join(cells)


def _extract_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    return "\n\n".join(_iter_docx_text(document)).strip()


def extract_text_from_file(content: bytes, filename: str) -> tuple[str, str]:
    normalized_name = filename.lower()
    if normalized_name.endswith(".txt"):
        return _extract_txt(content), "txt"
    if normalized_name.endswith(".pdf"):
        return _extract_pdf(content), "pdf"
    if normalized_name.endswith(".docx"):
        return _extract_docx(content), "docx"

    raise UnsupportedFileTypeError("รองรับเฉพาะไฟล์ .txt, .pdf และ .docx เท่านั้น")
